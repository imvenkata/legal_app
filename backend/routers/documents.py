from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Query
from pydantic import BaseModel
from typing import List, Dict, Any, Optional, Union
import os
from datetime import datetime
import aiofiles
from pathlib import Path
import json
import logging
import PyPDF2
import io
from PIL import Image
import base64
from dotenv import load_dotenv
import torch
import re
from llm_adapter.factory import LLMAdapterFactory
import docx
from pdf2image import convert_from_path
import tempfile

# Try to import ColPali models, but provide fallback if they're not available
try:
    from transformers import AutoTokenizer, AutoModel
    # Only try to import ColPali if transformers version is high enough
    COLPALI_AVAILABLE = False
    logging.info("Using standard transformer models as fallback")
except ImportError:
    COLPALI_AVAILABLE = False
    from transformers import AutoTokenizer, AutoModel
    logging.warning("Transformers library not available, using basic fallback implementation")

import asyncio

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize LLM adapter factory
llm_factory = LLMAdapterFactory()

# Get LLM configuration from environment variables
DEFAULT_LLM_PROVIDER = "openai" # Default if env var not set
DEFAULT_LLM_MODEL = "gpt-4"     # Default if env var not set
LLM_PROVIDER = os.getenv("LLM_PROVIDER", DEFAULT_LLM_PROVIDER).lower()
LLM_MODEL_NAME = os.getenv("LLM_MODEL_NAME", DEFAULT_LLM_MODEL)
# Add ColPali specific env vars
COLPALI_TOKENIZER_NAME = os.getenv("COLPALI_TOKENIZER_NAME") # Required if LLM_PROVIDER is colpali
DEVICE = os.getenv("DEVICE") # Optional device override (e.g., "cpu", "cuda:0")

logger.info(f"Using LLM Provider: {LLM_PROVIDER}, Model: {LLM_MODEL_NAME}")
if LLM_PROVIDER == "colpali":
    logger.info(f"ColPali Tokenizer: {COLPALI_TOKENIZER_NAME}")
if DEVICE:
    logger.info(f"Using Device: {DEVICE}")

router = APIRouter()

# Create uploads directory if it doesn't exist
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
logger.info(f"Upload directory: {UPLOAD_DIR.absolute()}")

# Initialize model and processor (lazy loading to save memory)
MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
model = None
processor = None

def get_model():
    global model, processor
    if model is None:
        logger.info(f"Initializing model: {MODEL_NAME}")
        device = "cuda:0" if torch.cuda.is_available() else "cpu"
        logger.info(f"Using device: {device}")
        
        try:
            # Load model with appropriate precision
            dtype = torch.bfloat16 if torch.cuda.is_available() else torch.float32
            
            # Use sentence-transformers for document embedding
            model = AutoModel.from_pretrained(MODEL_NAME).to(device).eval()
            processor = AutoTokenizer.from_pretrained(MODEL_NAME)
            logger.info("Sentence transformer model initialized successfully")
                
        except Exception as e:
            logger.error(f"Error initializing model: {str(e)}")
            raise e
            
    return model, processor

# Models
class DocumentBase(BaseModel):
    title: str
    description: Optional[str] = None

class DocumentCreate(DocumentBase):
    pass

class DocumentResponse(DocumentBase):
    id: str
    file_path: str
    file_type: str
    status: str
    created_at: str

class DocumentAnalysisRequest(BaseModel):
    model: str = "document-analyzer"

class DocumentChatRequest(BaseModel):
    query: str
    model: str = "document-analyzer"

class ChatMessage(BaseModel):
    role: str
    content: str

class DocumentChatResponse(BaseModel):
    document_id: str
    response: str
    messages: List[ChatMessage]

class Entity(BaseModel):
    name: str
    type: str
    mentions: List[str]

class DocumentAnalysisResponse(BaseModel):
    document_id: str
    summary: str
    key_points: List[str]
    risks: List[str]
    entities: List[Entity]
    model: str

# In-memory database for documents and chat histories
documents_db = {}
document_chat_histories = {}
document_embeddings = {}  # Store document embeddings for efficient retrieval

# Initialize LLM adapter (deferred initialization)
llm_adapter_instance = None

async def get_llm_adapter():
    """Initializes and returns the LLM adapter instance on demand."""
    global llm_adapter_instance
    if llm_adapter_instance is None:
        try:
            adapter_class = llm_factory.get_adapter_class(LLM_PROVIDER, LLM_MODEL_NAME)
            llm_adapter_instance = adapter_class() # Instantiate

            # Prepare initialization parameters
            model_params = {}
            if DEVICE:
                model_params["device"] = DEVICE
            
            # Call the specific initialize method
            if LLM_PROVIDER == "openai":
                api_key = os.getenv("OPENAI_API_KEY")
                if not api_key:
                     raise ValueError("OPENAI_API_KEY environment variable not set.")
                model_params["model"] = LLM_MODEL_NAME # Pass model name
                await llm_adapter_instance.initialize(api_key=api_key, model_params=model_params)
            
            elif LLM_PROVIDER == "colpali":
                 if not COLPALI_TOKENIZER_NAME:
                     raise ValueError("COLPALI_TOKENIZER_NAME environment variable required for ColPali.")
                 await llm_adapter_instance.initialize(model_name=LLM_MODEL_NAME, 
                                                       tokenizer_name=COLPALI_TOKENIZER_NAME, 
                                                       model_params=model_params)
            # Add initialization for other adapters (gemini, deepseek) here
            # elif LLM_PROVIDER == "gemini": ...
            # elif LLM_PROVIDER == "deepseek": ...
            else:
                 # Basic initialization if adapter supports it
                 logger.warning(f"Adapter for provider '{LLM_PROVIDER}' does not have explicit async initialization logic in get_llm_adapter. Ensure its __init__ is sufficient.")
                 # Or raise an error if initialization is always required

            logger.info(f"Successfully initialized LLM adapter for {LLM_PROVIDER}")
        except Exception as e:
            logger.error(f"Failed to initialize LLM adapter: {e}", exc_info=True)
            # Prevent further execution if adapter fails
            raise RuntimeError(f"LLM Adapter initialization failed: {e}") 
            
    return llm_adapter_instance

async def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from PDF file."""
    try:
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to extract text from PDF file: {str(e)}"
        )

async def extract_text_from_image(file_path: Path) -> str:
    """Placeholder for OCR functionality"""
    # In a real implementation, you would use an OCR library like pytesseract
    return "Image document - text extraction would require OCR processing."

def extract_entities(text: str) -> List[Dict[str, Any]]:
    """Extract entities using simple regex patterns"""
    entities = []
    
    # Find potential people names (capitalized words)
    person_pattern = r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b'
    people = re.findall(person_pattern, text)
    
    # Find potential organizations (uppercase words)
    org_pattern = r'\b[A-Z][A-Z]+\b'
    orgs = re.findall(org_pattern, text)
    
    # Add people entities
    for person in set(people):
        entities.append({
            "name": person,
            "type": "person",
            "mentions": ["Found in document"]
        })
    
    # Add organization entities
    for org in set(orgs):
        entities.append({
            "name": org,
            "type": "organization",
            "mentions": ["Found in document"]
        })
    
    return entities

async def extract_entities_with_vlm(document_path: Path, file_type: str) -> List[Dict[str, Any]]:
    """Extract entities page-by-page using a Vision Language Model (ColPali)."""
    all_entities = []
    processed_entity_keys = set()
    adapter = await get_llm_adapter()
    
    page_images: List[Image.Image] = []

    logger.info(f"Starting VLM entity extraction for {document_path}")

    if file_type == ".pdf":
        try:
            # Use a temporary directory for pdf2image if needed, or let it manage
            logger.info(f"Converting PDF to images: {document_path}")
            page_images = convert_from_path(document_path, dpi=200) # Adjust DPI as needed
            logger.info(f"Converted PDF to {len(page_images)} images.")
        except Exception as e:
             logger.error(f"Failed to convert PDF to images: {e}", exc_info=True)
             raise HTTPException(status_code=500, detail=f"PDF processing failed: {e}")
    elif file_type in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
        try:
             logger.info(f"Loading image file: {document_path}")
             img = Image.open(document_path)
             page_images = [img] # Treat single image as one page
        except Exception as e:
             logger.error(f"Failed to load image file: {e}", exc_info=True)
             raise HTTPException(status_code=400, detail=f"Image loading failed: {e}")
    # Add handling for DOCX to image conversion if a library/tool is available
    # elif file_type == ".docx":
    #     logger.warning("DOCX to image conversion not implemented for VLM processing.")
    #     # Fallback or raise error?
    else:
        logger.warning(f"VLM entity extraction not supported for file type: {file_type}")
        return [] # Cannot process non-image/PDF files with VLM page-by-page

    if not page_images:
        logger.warning(f"No images generated or loaded for VLM processing: {document_path}")
        return []

    # Process each page image
    for i, page_image in enumerate(page_images):
        logger.info(f"Processing page {i+1}/{len(page_images)} with VLM for entities.")
        try:
            # Call the adapter method responsible for extracting entities from an image
            page_entities = await adapter.extract_entities_from_page(page_image)
            
            # Deduplicate entities found on this page against overall list
            for entity in page_entities:
                 if isinstance(entity, dict) and "name" in entity and "type" in entity:
                     entity_key = (entity["name"].strip(), entity["type"].strip())
                     if entity_key not in processed_entity_keys:
                         all_entities.append({
                             "name": entity["name"].strip(),
                             "type": entity["type"].strip(),
                             "mentions": [entity.get("context", f"Found on page {i+1}")]
                         })
                         processed_entity_keys.add(entity_key)
                 else:
                      logger.warning(f"Skipping malformed entity from VLM on page {i+1}: {entity}")
                      
        except Exception as e:
            logger.error(f"Error processing page {i+1} with VLM adapter: {e}", exc_info=True)
            # Decide whether to continue with other pages or fail
            continue 
        finally:
            # Clean up the image object to save memory
            page_image.close()
            
    logger.info(f"VLM extraction completed. Found {len(all_entities)} unique entities.")
    return all_entities

# Keep the text-based LLM entity extraction for non-VLM providers
async def extract_entities_with_llm(text: str, chunk_size: int = 4000, overlap: int = 200) -> List[Dict[str, Any]]:
    """Extract entities from text using LLM with chunking."""
    all_formatted_entities = []
    processed_entities = set() # Keep track of unique entities (name, type)

    try:
        adapter = await get_llm_adapter() # Get initialized adapter
        logger.info(f"Using LLM adapter for {LLM_PROVIDER} - {LLM_MODEL_NAME}")

        # Split text into overlapping chunks
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start += chunk_size - overlap # Move window with overlap

        logger.info(f"Processing text in {len(chunks)} chunks for entity extraction.")

        for i, chunk in enumerate(chunks):
            logger.info(f"Processing chunk {i+1}/{len(chunks)}")
            # Create a refined prompt for entity extraction for each chunk
            prompt = f"""
            Extract all named entities from the following legal document text chunk.
            Focus on identifying specific names of people, organizations, locations, dates, and monetary values.
            For each entity, provide:
            1. The exact entity name as it appears in the text.
            2. The entity type (must be one of: PERSON, ORGANIZATION, LOCATION, DATE, MONEY, OTHER).
            3. A brief context (5-10 words) of how the entity is mentioned in this chunk.

            Format your response ONLY as a valid JSON array of objects. Each object must have "name", "type", and "context" keys.
            Example:
            [
                {{"name": "John Doe", "type": "PERSON", "context": "referred to as the primary signatory"}},
                {{"name": "Acme Corp", "type": "ORGANIZATION", "context": "mentioned as the contracting party"}},
                {{"name": "January 1, 2023", "type": "DATE", "context": "stated as the agreement effective date"}}
            ]

            If no entities are found in this chunk, return an empty JSON array [].

            Document text chunk:
            ---
            {chunk}
            ---
            """

            try:
                # Get response from LLM
                response = await adapter.generate_text(prompt, max_tokens=1000) # Adjust max_tokens if needed

                # Attempt to parse the JSON response
                try:
                    # Clean potential markdown code blocks
                    if response.strip().startswith("```json"):
                        response = response.strip()[7:-3].strip()
                    elif response.strip().startswith("```"):
                         response = response.strip()[3:-3].strip()

                    entities = json.loads(response)

                    # Validate structure and add unique entities
                    if isinstance(entities, list):
                        for entity in entities:
                            if isinstance(entity, dict) and "name" in entity and "type" in entity:
                                entity_key = (entity["name"].strip(), entity["type"].strip())
                                # Add entity if it hasn't been processed already
                                if entity_key not in processed_entities:
                                     all_formatted_entities.append({
                                        "name": entity["name"].strip(),
                                        "type": entity["type"].strip(),
                                        # Use context as the first mention
                                        "mentions": [entity.get("context", "Context not provided").strip()]
                                    })
                                     processed_entities.add(entity_key)
                            else:
                                logger.warning(f"Skipping malformed entity object in chunk {i+1}: {entity}")
                    else:
                         logger.warning(f"LLM response for chunk {i+1} was not a JSON list: {response[:100]}...")

                except json.JSONDecodeError as json_err:
                    logger.error(f"Failed to parse LLM response for chunk {i+1} as JSON: {json_err}")
                    logger.error(f"LLM Raw Response (chunk {i+1}): {response}")
                    # Optionally, you could try a fallback here for this chunk, or just skip it.
                    continue # Move to the next chunk

            except Exception as llm_err:
                 logger.error(f"Error during LLM call for chunk {i+1}: {str(llm_err)}")
                 continue # Move to the next chunk

        if not all_formatted_entities:
             logger.warning("LLM entity extraction returned no entities. Falling back to basic extraction.")
             return extract_entities(text) # Fallback if LLM fails completely

        logger.info(f"Successfully extracted {len(all_formatted_entities)} unique entities using LLM.")
        return all_formatted_entities

    except Exception as e:
        logger.error(f"Error in LLM entity extraction process: {str(e)}")
        logger.info("Falling back to basic regex entity extraction.")
        # Fallback to basic entity extraction if main process fails
        return extract_entities(text)

def extract_key_points(text: str, max_points: int = 5) -> List[str]:
    """Extract key points from text"""
    # Split text into sentences
    sentences = re.split(r'[.!?]', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
    
    # Find sentences with important keywords
    important_keywords = ['must', 'shall', 'required', 'important', 'agreement', 'contract', 'legal', 'rights', 'obligations']
    key_sentences = []
    
    for sentence in sentences:
        if any(keyword in sentence.lower() for keyword in important_keywords):
            key_sentences.append(sentence)
    
    # If we don't have enough key sentences, add some of the longest sentences
    if len(key_sentences) < max_points:
        remaining = max_points - len(key_sentences)
        sorted_by_length = sorted(sentences, key=len, reverse=True)
        for sentence in sorted_by_length:
            if sentence not in key_sentences:
                key_sentences.append(sentence)
                remaining -= 1
                if remaining <= 0:
                    break
    
    return key_sentences[:max_points]

def identify_risks(text: str) -> List[str]:
    """Identify potential risks in the text"""
    risk_keywords = [
        'risk', 'liability', 'termination', 'penalty', 'breach', 'sue', 'lawsuit',
        'damage', 'claim', 'litigation', 'conflict', 'dispute', 'deadline',
        'fail', 'default', 'violation', 'responsible', 'obligation'
    ]
    
    risks = []
    sentences = re.split(r'[.!?]', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
    
    for sentence in sentences:
        if any(keyword in sentence.lower() for keyword in risk_keywords):
            risks.append(sentence)
    
    # Limit to 3 risks to avoid overwhelming
    return risks[:3]

async def analyze_document(file_path: Path) -> Dict[str, Any]:
    """
    Analyze document using available models (Text-based or VLM based on config).
    """
    adapter = await get_llm_adapter() # Ensure adapter is initialized
    file_type = file_path.suffix.lower()
    text = "" # Initialize text variable
    entity_list = []
    summary = ""
    key_points = []
    risks = []

    try:
        # --- VLM Path (ColPali) --- 
        if LLM_PROVIDER == 'colpali':
             logger.info("Using VLM (ColPali) analysis path.")
             # For VLMs, primary analysis might be page-based.
             # We might extract entities via VLM first.
             entity_list = await extract_entities_with_vlm(file_path, file_type)
             
             # Text extraction might still be needed for summary/key points/risks
             # if the VLM isn't used for those, or as fallback.
             logger.info("Extracting text for non-entity analysis (summary, key points, risks)...")
             # TODO: Consolidate text extraction logic? Refactor needed.
             if file_type == ".pdf":
                 text = await extract_text_from_pdf(file_path)
             elif file_type == ".docx":
                 try:
                     doc = docx.Document(file_path)
                     full_text = []
                     for para in doc.paragraphs:
                         full_text.append(para.text)
                     text = '\n'.join(full_text)
                 except Exception as e:
                      logger.error(f"Error extracting text from DOCX: {str(e)}")
                      raise HTTPException(
                          status_code=status.HTTP_400_BAD_REQUEST,
                          detail=f"Failed to extract text from DOCX file: {str(e)}"
                      )
             elif file_type == ".txt":
                  # ... (txt extraction logic) ...
                  logger.warning("Using placeholder text for text-based analysis.")
                  text = "Text content - requires VLM OCR for text-based analysis."
             elif file_type in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
                  # Maybe use VLM for OCR here? 
                  # text = await adapter.process_page(Image.open(file_path), "Extract all text from this image.")
                  logger.warning("Using placeholder text for image analysis.")
                  text = "Image content - requires VLM OCR for text-based analysis."
             else:
                  logger.warning(f"Cannot extract text for file type {file_type}")
                  text = ""
             
             # Generate summary, key points, risks using the extracted text (or maybe another VLM call?)
             if text:
                 sentences = re.split(r'[.!?]', text)
                 sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
                 summary = ". ".join(sentences[:3]) + "." if sentences else text[:500]
                 key_points = extract_key_points(text)
                 risks = identify_risks(text)
             else:
                  logger.warning("No text extracted, skipping text-based analysis (summary, key points, risks).")
                  summary = "Analysis based on VLM entities only."
                  key_points = ["Refer to extracted entities."]
                  risks = ["Text not analyzed."]

        # --- Text-Based Path (OpenAI, DeepSeek, etc.) --- 
        else:
            logger.info("Using text-based analysis path.")
            # 1. Extract Text first
            if file_type == ".pdf":
                text = await extract_text_from_pdf(file_path)
            elif file_type == ".docx":
                 try:
                     doc = docx.Document(file_path)
                     full_text = []
                     for para in doc.paragraphs:
                         full_text.append(para.text)
                     text = '\n'.join(full_text)
                 except Exception as e:
                      logger.error(f"Error extracting text from DOCX: {str(e)}")
                      raise HTTPException(
                          status_code=status.HTTP_400_BAD_REQUEST,
                          detail=f"Failed to extract text from DOCX file: {str(e)}"
                      )
            elif file_type == ".txt":
                 # ... (existing txt handling) ...
                 logger.warning("Using placeholder text for text-based analysis.")
                 text = "Text content - requires VLM OCR for text-based analysis."
            elif file_type in [".jpg", ".jpeg", ".png"]:
                 text = await extract_text_from_image(file_path)
                 logger.warning("Image text extraction is currently a placeholder.")
            else:
                 logger.warning(f"Unsupported file type for text extraction: {file_type}. Analysis may be incomplete.")
                 text = f"Cannot extract text from {file_type} files automatically."

            if not text and file_type not in [".jpg", ".jpeg", ".png"]:
                 logger.error(f"Extracted text is empty for file: {file_path}")
                 text = ""
            
            logger.info(f"Successfully extracted text (length: {len(text)}) from {file_path}")

            # 2. Perform analysis on the extracted text
            if text:
                sentences = re.split(r'[.!?]', text)
                sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
                summary = ". ".join(sentences[:3]) + "." if sentences else text[:500]
                key_points = extract_key_points(text)
                risks = identify_risks(text)
                # Extract entities using the appropriate text-based LLM function
                entity_list = await extract_entities_with_llm(text)
            else:
                 logger.warning("Skipping analysis due to empty extracted text.")
                 summary = "Failed to extract text."
                 entity_list = []
                 key_points = []
                 risks = []

        # --- Consolidate Results --- 
        result = {
            "summary": summary,
            "key_points": key_points or ["Not extracted"],
            "risks": risks or ["No specific risks identified or text not analyzed"],
            "entities": entity_list or [] # Ensure it's always a list
        }
        
        return result
            
    except Exception as e:
        logger.error(f"Error in document analysis: {str(e)}", exc_info=True)
        # Ensure a response is still sent in case of error
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Analysis failed: {str(e)}"
        )

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    title: str = Form(...),
    description: Optional[str] = Form(None)
):
    try:
        logger.info(f"Starting file upload: {file.filename}")
        
        # Validate file type first
        file_type = file.filename.split(".")[-1].lower() if "." in file.filename else ""
        allowed_types = ["pdf", "doc", "docx", "txt", "jpg", "jpeg", "png"]
        
        if file_type not in allowed_types:
            logger.error(f"Invalid file type: {file_type}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type not allowed. Allowed types: {', '.join(allowed_types)}"
            )
        
        # Create a unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{file.filename}"
        file_path = UPLOAD_DIR / safe_filename
        logger.info(f"Saving file to: {file_path}")
        
        # Save the file with progress logging
        try:
            content = await file.read()
            file_size = len(content)
            logger.info(f"File size: {file_size} bytes")
            
            async with aiofiles.open(file_path, 'wb') as out_file:
                await out_file.write(content)
            logger.info("File saved successfully")
        except Exception as e:
            logger.error(f"Error saving file: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to save file: {str(e)}"
            )
        
        # Create document record
        document_id = str(len(documents_db) + 1)
        document = {
            "id": document_id,
            "title": title,
            "description": description,
            "file_path": str(file_path),
            "file_type": file_type,
            "status": "uploaded",
            "created_at": datetime.now().isoformat()
        }
        
        documents_db[document_id] = document
        
        # Initialize empty chat history for this document
        document_chat_histories[document_id] = []
        
        logger.info(f"Document record created: {document_id}")
        
        return document
    except HTTPException as he:
        logger.error(f"HTTP Exception during upload: {str(he)}")
        raise he
    except Exception as e:
        logger.error(f"Unexpected error during upload: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Upload failed: {str(e)}"
        )

@router.post("/{document_id}/analyze", response_model=DocumentAnalysisResponse)
async def analyze_document_endpoint(document_id: str, request: DocumentAnalysisRequest):
    if document_id not in documents_db:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    try:
        document = documents_db[document_id]
        file_path = Path(document["file_path"])
        
        if not file_path.exists():
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document file not found"
            )
        
        # Analyze with available model (handles text or VLM path internally)
        analysis_result = await analyze_document(file_path)
        
        # Update document status
        documents_db[document_id]["status"] = "analyzed"
        
        # Store the full analysis for future chats
        documents_db[document_id]["analysis"] = analysis_result
        
        # Return the formatted response
        return DocumentAnalysisResponse(
            document_id=document_id,
            model=LLM_MODEL_NAME, # Report the configured model
            **analysis_result # Unpack the dict from analyze_document
        )
    except Exception as e:
        logger.error(f"Error analyzing document: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e)
        )

@router.post("/{document_id}/chat", response_model=DocumentChatResponse)
async def chat_with_document(document_id: str, request: DocumentChatRequest):
    if document_id not in documents_db:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    document = documents_db[document_id]
    file_path = Path(document["file_path"])
    adapter = await get_llm_adapter() # Get adapter

    try:
        chat_history = document_chat_histories.get(document_id, [])
        file_type = document["file_type"]
        response_text = ""

        # --- VLM Chat Path (Placeholder) --- 
        if LLM_PROVIDER == 'colpali':
            # TODO: Implement VLM-based chat
            # - Could try to find relevant page images based on query embedding?
            # - Could ask VLM a question about the document using page images as context?
            logger.warning("Chat endpoint not fully implemented for ColPali VLM yet.")
            response_text = f"VLM chat is under development. Your query was: {request.query}"
        
        # --- Text-Based Chat Path --- 
        else:
            # Extract text (similar logic as in analyze_document text path)
            # ... (Add text extraction logic here for different file types) ...
            text = "" 
            if file_type == ".pdf":
                 text = await extract_text_from_pdf(file_path)
            # Add other types (docx, txt...)
            else:
                 async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                     text = await f.read()

            if not text:
                 response_text = "Could not read document text to answer query."
            else:
                 # Existing simple keyword matching logic
                 query = request.query.lower()
                 # ... (rest of the keyword matching logic) ...
                 # Maybe enhance this with an LLM call using the text context:
                 # prompt = f"Document context:\n{text[:8000]}\n\nChat History:\n{chat_history}\n\nUser Query: {request.query}\n\nAnswer the user query based *only* on the provided document context and chat history."
                 # response_text = await adapter.generate_text(prompt)

        # Add messages to history
        chat_history.append({"role": "user", "content": request.query})
        chat_history.append({"role": "assistant", "content": response_text})
        document_chat_histories[document_id] = chat_history
        
        return DocumentChatResponse(
            document_id=document_id,
            response=response_text,
            messages=chat_history
        )
        
    except Exception as e:
         logger.error(f"Error in document chat: {str(e)}", exc_info=True)
         raise HTTPException(
             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
             detail=f"Chat failed: {str(e)}"
         )

@router.get("/{document_id}/chat_history", response_model=List[ChatMessage])
async def get_chat_history(document_id: str):
    if document_id not in documents_db:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    
    # Return the chat history for this document
    return document_chat_histories.get(document_id, [])

@router.get("/", response_model=List[DocumentResponse])
async def get_documents():
    return list(documents_db.values())

@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(document_id: str):
    if document_id not in documents_db:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
    return documents_db[document_id]
